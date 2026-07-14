"""Extração de texto de anexos de contrato: PDF (com fallback OCR), imagens, DOCX e XLSX."""

import io
import logging
import re

from app.config import get_settings

logger = logging.getLogger(__name__)

# Página de PDF com menos texto extraível que isso é tratada como digitalizada → OCR.
_MIN_CHARS_PAGINA_PDF = 30

_EXTENSOES_IMAGEM = {"png", "jpg", "jpeg"}


class ExtracaoError(Exception):
    """Falha ao extrair texto de um documento."""


def extrair_texto(conteudo: bytes, nome_arquivo: str | None) -> str:
    """Extrai texto do documento conforme extensão (fallback: magic bytes)."""
    ext = _detectar_tipo(conteudo, nome_arquivo)

    if ext == "pdf":
        return _extrair_pdf(conteudo)
    if ext in _EXTENSOES_IMAGEM:
        return _ocr_imagem(conteudo)
    if ext == "docx":
        return _extrair_docx(conteudo)
    if ext == "xlsx":
        return _extrair_xlsx(conteudo)

    # Fallback: tenta decodificar como texto plano.
    try:
        return conteudo.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return conteudo.decode("latin-1")
        except Exception as e:  # pragma: no cover
            raise ExtracaoError(f"Tipo de arquivo não suportado: {nome_arquivo!r}") from e


def _detectar_tipo(conteudo: bytes, nome_arquivo: str | None) -> str:
    if nome_arquivo and "." in nome_arquivo:
        return nome_arquivo.rsplit(".", 1)[-1].lower()
    # Magic bytes
    if conteudo.startswith(b"%PDF"):
        return "pdf"
    if conteudo.startswith(b"\x89PNG"):
        return "png"
    if conteudo.startswith(b"\xff\xd8\xff"):
        return "jpg"
    if conteudo.startswith(b"PK"):
        # Office Open XML (docx/xlsx) — sem extensão, tenta docx.
        return "docx"
    return ""


def _extrair_pdf(conteudo: bytes) -> str:
    # pdfium (engine do Chrome) preserva espaços entre palavras — pypdf concatena
    # palavras em alguns PDFs (ex.: gerados pelo PAE/sistemas.pa.gov.br).
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(conteudo)
    paginas: list[str] = []
    paginas_para_ocr: list[int] = []
    try:
        for i in range(len(pdf)):
            texto = pdf[i].get_textpage().get_text_bounded().strip()
            if len(texto) < _MIN_CHARS_PAGINA_PDF:
                paginas_para_ocr.append(i)
                paginas.append("")
            else:
                paginas.append(texto)

        if paginas_para_ocr:
            logger.info(
                "PDF com %s página(s) sem texto extraível; aplicando OCR", len(paginas_para_ocr)
            )
            try:
                for i in paginas_para_ocr:
                    paginas[i] = _ocr_pagina_pdf(pdf, i)
            except Exception:
                # OCR é best-effort em PDF misto: não perde as páginas com texto nativo.
                # PDF 100% digitalizado sem OCR resulta em texto vazio → erro na ingestão.
                logger.exception(
                    "OCR falhou; seguindo apenas com o texto nativo das demais páginas"
                )
    finally:
        pdf.close()

    paginas = _remover_linhas_repetidas(paginas)
    paginas = [re.sub(r"\s+", " ", p).strip() for p in paginas]
    # Une páginas com espaço: frase que quebra na virada de página vira uma frase só
    # (o rodapé/cabeçalho repetido que ficava no meio já foi removido acima).
    return " ".join(p for p in paginas if p)


# Linha presente em >=30% das páginas (mín. 3) é cabeçalho/rodapé (timbre, carimbo
# de assinatura eletrônica, protocolo) — ruído que polui os embeddings dos chunks.
_FRACAO_LINHA_REPETIDA = 0.3
_MIN_PAGINAS_LINHA_REPETIDA = 3


def _remover_linhas_repetidas(paginas: list[str]) -> list[str]:
    if len(paginas) < _MIN_PAGINAS_LINHA_REPETIDA:
        return paginas

    frequencia: dict[str, int] = {}
    for pagina in paginas:
        for linha in {ln.strip() for ln in pagina.splitlines() if ln.strip()}:
            frequencia[linha] = frequencia.get(linha, 0) + 1

    limite = max(_MIN_PAGINAS_LINHA_REPETIDA, int(len(paginas) * _FRACAO_LINHA_REPETIDA))
    repetidas = {ln for ln, freq in frequencia.items() if freq >= limite}
    if not repetidas:
        return paginas

    logger.info("Removendo %s linha(s) de cabeçalho/rodapé repetidas", len(repetidas))
    return [
        "\n".join(ln for ln in pagina.splitlines() if ln.strip() not in repetidas)
        for pagina in paginas
    ]


def _ocr_pagina_pdf(pdf, indice: int) -> str:
    import pytesseract

    settings = get_settings()
    bitmap = pdf[indice].render(scale=2.0)
    return pytesseract.image_to_string(bitmap.to_pil(), lang=settings.ocr_lang)


def _ocr_imagem(conteudo: bytes) -> str:
    import pytesseract
    from PIL import Image

    settings = get_settings()
    imagem = Image.open(io.BytesIO(conteudo))
    return pytesseract.image_to_string(imagem, lang=settings.ocr_lang)


def _extrair_docx(conteudo: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(conteudo))
    partes: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
            if celulas:
                partes.append(" | ".join(celulas))
    return "\n".join(partes)


def _extrair_xlsx(conteudo: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(conteudo), read_only=True, data_only=True)
    partes: list[str] = []
    try:
        for sheet in wb.worksheets:
            partes.append(f"[Planilha: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                valores = [str(v) for v in row if v is not None and str(v).strip()]
                if valores:
                    partes.append(" | ".join(valores))
    finally:
        wb.close()
    return "\n".join(partes)
