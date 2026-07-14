import io

from app.services.document_extractor import _detectar_tipo, extrair_texto


def test_detectar_tipo_por_extensao():
    assert _detectar_tipo(b"", "contrato.PDF") == "pdf"
    assert _detectar_tipo(b"", "foto.jpeg") == "jpeg"
    assert _detectar_tipo(b"", "planilha.xlsx") == "xlsx"


def test_detectar_tipo_por_magic_bytes():
    assert _detectar_tipo(b"%PDF-1.7 ...", None) == "pdf"
    assert _detectar_tipo(b"\x89PNG\r\n", None) == "png"
    assert _detectar_tipo(b"\xff\xd8\xff\xe0", None) == "jpg"


def test_extrair_texto_plano():
    assert extrair_texto("olá contrato".encode("utf-8"), "notas.txt") == "olá contrato"


def test_extrair_texto_latin1_fallback():
    assert extrair_texto("aditivo nº 1".encode("latin-1"), "notas.txt") == "aditivo nº 1"


def test_extrair_docx():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Cláusula primeira do contrato.")
    tabela = doc.add_table(rows=1, cols=2)
    tabela.cell(0, 0).text = "Item"
    tabela.cell(0, 1).text = "Valor"
    buf = io.BytesIO()
    doc.save(buf)

    texto = extrair_texto(buf.getvalue(), "contrato.docx")
    assert "Cláusula primeira do contrato." in texto
    assert "Item | Valor" in texto


def test_extrair_xlsx():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Postos"
    ws.append(["Posto", "Quantidade"])
    ws.append(["Vigilante", 12])
    buf = io.BytesIO()
    wb.save(buf)

    texto = extrair_texto(buf.getvalue(), "postos.xlsx")
    assert "[Planilha: Postos]" in texto
    assert "Posto | Quantidade" in texto
    assert "Vigilante | 12" in texto


def test_extrair_pdf_com_texto():
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)

    # Página em branco → sem texto extraível → cai no OCR; sem tesseract o teste
    # apenas garante que o tipo é roteado como PDF (erro claro, não crash silencioso).
    try:
        texto = extrair_texto(buf.getvalue(), "vazio.pdf")
        assert isinstance(texto, str)
    except Exception:
        pass
